from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parent / "CRS365-AI-Automation-Blueprint-Template.docx"
NAVY, NAVY2, TEAL, GOLD = "0F2744", "1A3A5C", "0D5F5F", "B88718"
INK, MUTED, LINE, SOFT = "1F3549", "5C7082", "DCE4E9", "F3F7F7"
WHITE, NOTE, GREEN = "FFFFFF", "FFF7E3", "147A50"
WIDTH, INDENT = 9360, 120

def color(value):
    return RGBColor.from_string(value)

def font(run, size=11, shade=INK, bold=None, italic=None):
    run.font.name = "Calibri"
    props = run._element.get_or_add_rPr()
    props.rFonts.set(qn("w:ascii"), "Calibri")
    props.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color(shade)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        props.append(node)
    node.set(qn("w:fill"), fill)

def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def borders(table, shade_value=LINE, size=6):
    props = table._tbl.tblPr
    group = props.find(qn("w:tblBorders"))
    if group is None:
        group = OxmlElement("w:tblBorders")
        props.append(group)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = group.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            group.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), shade_value)

def geometry(table, widths, indent=INDENT):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    props = table._tbl.tblPr
    width_node = props.find(qn("w:tblW"))
    if width_node is None:
        width_node = OxmlElement("w:tblW")
        props.append(width_node)
    width_node.set(qn("w:w"), str(sum(widths)))
    width_node.set(qn("w:type"), "dxa")
    indent_node = props.find(qn("w:tblInd"))
    if indent_node is None:
        indent_node = OxmlElement("w:tblInd")
        props.append(indent_node)
    indent_node.set(qn("w:w"), str(indent))
    indent_node.set(qn("w:type"), "dxa")
    layout = props.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        props.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(value))
        grid.append(node)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            props = cell._tc.get_or_add_tcPr()
            cell_width = props.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                props.append(cell_width)
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def cell_text(cell, text, bold=False, shade_value=INK, size=9.7):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(str(text))
    font(run, size=size, shade=shade_value, bold=bold)

def table(doc, headers, rows, widths, header_fill=NAVY2):
    item = doc.add_table(rows=1, cols=len(headers))
    geometry(item, widths)
    borders(item)
    header = item.rows[0]
    header_props = header._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_props.append(repeat)
    for index, heading in enumerate(headers):
        shade(header.cells[index], header_fill)
        cell_text(header.cells[index], heading, True, WHITE, 9.2)
    for row_index, values in enumerate(rows):
        row = item.add_row()
        for index, value in enumerate(values):
            cell_text(row.cells[index], value)
            if row_index % 2:
                shade(row.cells[index], "F8FAFB")
    geometry(item, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return item

def body(doc, text, bold_start=None, size=11, shade_value=INK, italic=False, after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_start and text.startswith(bold_start):
        first = paragraph.add_run(bold_start)
        font(first, size=size, shade=shade_value, bold=True)
        second = paragraph.add_run(text[len(bold_start):])
        font(second, size=size, shade=shade_value, italic=italic)
    else:
        run = paragraph.add_run(text)
        font(run, size=size, shade=shade_value, italic=italic)
    return paragraph

def list_item(doc, text, numbered=False, bold_start=None):
    style = "List Number" if numbered else "List Bullet"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_start and text.startswith(bold_start):
        first = paragraph.add_run(bold_start)
        font(first, size=10.7, bold=True)
        second = paragraph.add_run(text[len(bold_start):])
        font(second, size=10.7)
    else:
        run = paragraph.add_run(text)
        font(run, size=10.7)
    return paragraph

def callout(doc, label, text, fill="DDEEEE", accent=TEAL):
    item = doc.add_table(rows=1, cols=1)
    geometry(item, [WIDTH])
    borders(item, accent, 8)
    shade(item.cell(0, 0), fill)
    cell = item.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    first = p.add_run(label + ": ")
    font(first, size=10.4, shade=accent, bold=True)
    second = p.add_run(text)
    font(second, size=10.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def guidance(doc, text):
    callout(doc, "TEMPLATE GUIDANCE", text, NOTE, "B88718")

def section_title(doc, number, label, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(number + "  " + label.upper())
    font(run, size=9.4, shade=TEAL, bold=True)
    h = doc.add_paragraph(style="Heading 1")
    h.paragraph_format.space_before = Pt(0)
    run = h.add_run(title)
    font(run, size=16, shade=NAVY2, bold=True)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def heading(doc, text, level=2):
    p = doc.add_paragraph(style="Heading " + str(level))
    run = p.add_run(text)
    font(run, size=13 if level == 2 else 11.5, shade=TEAL if level == 2 else NAVY2, bold=True)
    return p

def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    font(run, size=8.5, shade=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    shown = OxmlElement("w:t")
    shown.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, shown, end])

def configure(doc):
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.82)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.42)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = color(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, shade_value, before, after in (
        ("Heading 1", 16, NAVY2, 16, 8),
        ("Heading 2", 13, TEAL, 12, 6),
        ("Heading 3", 11.5, NAVY2, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color(shade_value)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.7)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.1
    hp = section.header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("CRS365  |  AI Automation Blueprint")
    font(run, size=8.5, shade=MUTED, bold=True)
    page_number(section.footer.paragraphs[0])

def main():
    doc = Document()
    configure(doc)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(8)
    font(p.add_run("CRS365"), size=13, shade=TEAL, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    font(p.add_run("AI Automation\nBlueprint"), size=30, shade=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    font(p.add_run("A decision-ready plan for one priority workflow."), size=15, shade=MUTED)

    cover = doc.add_table(rows=4, cols=2)
    rows = [
        ("PREPARED FOR", "Your Business"),
        ("PRIORITY WORKFLOW", "Priority Workflow"),
        ("PREPARED BY", "CRS365 Team"),
        ("ENGAGEMENT", "$995 Fixed-Scope Blueprint"),
    ]
    for index, values in enumerate(rows):
        cell_text(cover.cell(index, 0), values[0], True, TEAL, 8.7)
        cell_text(cover.cell(index, 1), values[1], index < 2, INK, 10.8)
        shade(cover.cell(index, 0), SOFT)
    geometry(cover, [2300, 7060])
    borders(cover)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    callout(doc, "Blueprint outcome", "A documented current state, future-state workflow, tool and data plan, human controls, success measures, and a sequenced 30-day rollout.")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    font(p.add_run("Optimize. Automate. Operate."), size=10, shade=GOLD, bold=True)

    page_break(doc)

    # 01 Executive summary
    section_title(doc, "01", "Executive Decision Summary", "The recommendation in one page")
    guidance(doc, "Replace all guidance and generic text before delivery. State the recommendation, why it matters now, and the first decision the client should make.")
    body(doc, "Recommendation: Define the future state in three to five sentences. Name the workflow, minimum useful automation, human approvals, and recommended implementation path.", "Recommendation:")
    body(doc, "Why now: Connect the recommendation to the client's stated friction, operating priority, and available readiness evidence.", "Why now:")
    body(doc, "First decision: State the specific approval or preparation decision required during the next seven days.", "First decision:")
    table(doc, ["Decision element", "Blueprint conclusion"], [
        ("Business objective", "State the measurable operational outcome."),
        ("Workflow", "Name the bounded process covered by this Blueprint."),
        ("Owner", "Name the person accountable for the future workflow."),
        ("Success measure", "Name the primary metric and current baseline."),
        ("Recommended path", "Proceed, prepare first, pilot with conditions, or do not automate."),
    ], [2700, 6660])
    callout(doc, "Executive takeaway", "Use one sentence the client can repeat internally without additional explanation.", "E8F2ED", GREEN)

    page_break(doc)

    # 02 Scope
    section_title(doc, "02", "Engagement Scope", "The workflow boundary and business objective")
    guidance(doc, "Keep the engagement to one workflow and no more than five core applications. Record assumptions that could change the recommendation.")
    table(doc, ["Scope item", "Definition"], [
        ("Trigger", "The event that starts the workflow."),
        ("Completed output", "The observable result that marks the workflow complete."),
        ("Business owner", "The person accountable for performance and exceptions."),
        ("In-scope roles", "People or teams that perform, approve, or receive the work."),
        ("In-scope systems", "Up to five applications reviewed in depth."),
        ("Primary objective", "The operational improvement the workflow should produce."),
    ], [2500, 6860])
    heading(doc, "Included")
    for item in (
        "Current-state workflow and friction analysis",
        "Readiness scorecard and future-state workflow",
        "Tool, data, integration, and control recommendations",
        "30-day roadmap, metrics, ownership, and findings review",
    ):
        list_item(doc, item)
    heading(doc, "Not included")
    for item in (
        "Software configuration, custom development, or data migration",
        "Vendor procurement, legal advice, or cybersecurity testing",
        "Additional workflows or broad organization-wide transformation",
        "Guaranteed savings, revenue, adoption, or performance outcomes",
    ):
        list_item(doc, item)
    body(doc, "Key assumptions: List the facts accepted as true for this analysis and identify which recommendation changes if an assumption is incorrect.", "Key assumptions:")

    page_break(doc)

    # 03 Current state
    section_title(doc, "03", "Current-State Workflow", "How the work operates today")
    guidance(doc, "Use concise stage names. Include evidence from the kickoff, SOPs, redacted examples, reports, and system observations.")
    table(doc, ["#", "Stage", "Owner / system", "Current action", "Friction or risk"], [
        ("1", "Trigger", "Role / system", "Describe how work begins.", "Delay, inconsistency, or missing input."),
        ("2", "Intake", "Role / system", "Describe collection and validation.", "Manual entry, duplicate data, or rework."),
        ("3", "Decision", "Role / system", "Describe judgment and approval.", "Unclear rules, bottleneck, or control gap."),
        ("4", "Handoff", "Role / system", "Describe routing and notification.", "Missed follow-up or disconnected tools."),
        ("5", "Completion", "Role / system", "Describe output and recordkeeping.", "Poor visibility, reporting, or ownership."),
    ], [450, 1250, 1850, 2850, 2960])
    heading(doc, "Observed friction")
    for label, text in (
        ("Frequency: ", "How often the problem occurs."),
        ("Impact: ", "Time, rework, delay, capacity, customer, or control impact."),
        ("Evidence: ", "What was observed or provided."),
        ("Root cause: ", "Why the friction exists rather than only where it appears."),
    ):
        list_item(doc, label + text, bold_start=label)
    heading(doc, "Baseline")
    table(doc, ["Measure", "Current baseline", "Source", "Confidence"], [
        ("Cycle time", "Record current time.", "System, estimate, or sample.", "High / medium / low"),
        ("Staff effort", "Record hours per cycle or month.", "Timesheet, estimate, or sample.", "High / medium / low"),
        ("Error / rework", "Record frequency or rate.", "System, issue log, or estimate.", "High / medium / low"),
        ("Response / completion", "Record service level or completion rate.", "Dashboard, CRM, or estimate.", "High / medium / low"),
    ], [2200, 2200, 2960, 2000])

    page_break(doc)

    # 04 Readiness
    section_title(doc, "04", "Readiness Scorecard", "What must be true before automation")
    guidance(doc, "Score each dimension from 1 to 5 using evidence. A low score is a preparation requirement, not a sales objection.")
    table(doc, ["Dimension", "Score", "Evidence", "Implementation implication"], [
        ("Process clarity", "1-5", "Rules, sequence, exceptions, and SOP evidence.", "What must be documented or simplified."),
        ("Data readiness", "1-5", "Completeness, consistency, access, and system of record.", "What must be cleaned or governed."),
        ("Ownership", "1-5", "Named owner, approvers, and exception responsibility.", "Who must own the future state."),
        ("Integration", "1-5", "APIs, exports, webhooks, permissions, and constraints.", "What can be connected and how."),
        ("Governance", "1-5", "Sensitive data, approvals, audit, retention, and access.", "Controls required before launch."),
        ("Adoption", "1-5", "User readiness, technical comfort, and change capacity.", "Training and rollout approach."),
    ], [1850, 850, 3100, 3560])
    callout(doc, "Scoring interpretation", "1 = undefined or unavailable, 3 = workable with preparation, 5 = clearly defined and ready to support a controlled pilot.")
    heading(doc, "Readiness conclusion")
    body(doc, "State whether the workflow is ready to pilot, ready with conditions, requires preparation, or should not be automated yet. Name the two most important readiness actions.")

    page_break(doc)

    # 05 Future state
    section_title(doc, "05", "Future-State Workflow", "How the improved process should operate")
    guidance(doc, "Design the minimum useful future state. Do not automate ambiguity. Keep judgment, approvals, and exceptions visible.")
    table(doc, ["#", "Future-stage action", "Automation role", "Human control", "Output"], [
        ("1", "Trigger and capture", "Detect event and create record.", "Review invalid or incomplete inputs.", "Validated work item"),
        ("2", "Enrich and route", "Add approved context and assign owner.", "Approve sensitive or ambiguous routing.", "Assigned work"),
        ("3", "Prepare decision", "Summarize, calculate, or draft using approved data.", "Make or approve material decision.", "Approved decision"),
        ("4", "Execute handoff", "Update systems, notify stakeholders, and create tasks.", "Resolve exceptions and failed actions.", "Completed handoff"),
        ("5", "Measure and learn", "Record status, timing, exceptions, and results.", "Review weekly and approve changes.", "Operational visibility"),
    ], [450, 2250, 2400, 2400, 1860])
    heading(doc, "Design principles")
    for item in (
        "Use one system of record for workflow status.",
        "Automate repeatable rules; escalate ambiguity.",
        "Require human approval for material, sensitive, or irreversible decisions.",
        "Make failures and exceptions visible to a named owner.",
        "Log enough information to audit outcomes and improve the workflow.",
    ):
        list_item(doc, item)
    page_break(doc)

    # 06 Tool and data
    section_title(doc, "06", "Tool and Data Plan", "What each system should own")
    guidance(doc, "Recommend roles before products. Prefer the current stack when it can meet the requirement without creating unacceptable risk or maintenance.")
    table(doc, ["System / category", "Primary role", "Data owned or exchanged", "Recommendation"], [
        ("System of record", "Authoritative workflow status and core records.", "Name records, identifiers, and required fields.", "Keep, configure, consolidate, or replace."),
        ("Automation layer", "Rules, routing, notifications, and data movement.", "Name triggers, actions, and error logs.", "Recommended platform role."),
        ("AI layer", "Summarization, classification, drafting, or extraction.", "Name approved inputs and prohibited data.", "Model, access, and review requirement."),
        ("Collaboration", "Tasks, communication, approvals, and handoffs.", "Name messages, tasks, and ownership fields.", "Channel and operating rule."),
        ("Reporting", "Baseline, exceptions, adoption, and outcome metrics.", "Name measures and source systems.", "Dashboard and review cadence."),
    ], [1900, 2250, 2500, 2710])
    heading(doc, "Purchase decision")
    body(doc, "State which existing tools can remain, which capabilities are missing, and whether a new subscription is necessary. Include expected licensing assumptions without presenting vendor pricing as permanent.")
    heading(doc, "Data movement")
    table(doc, ["From", "To", "Data", "Method", "Frequency / trigger", "Owner"], [
        ("Source system", "Destination system", "Minimum required fields", "Native integration, API, webhook, or export", "Event or schedule", "Role"),
        ("Source system", "Destination system", "Minimum required fields", "Native integration, API, webhook, or export", "Event or schedule", "Role"),
    ], [1250, 1250, 1850, 2050, 1760, 1200])

    page_break(doc)

    # 07 Controls
    section_title(doc, "07", "Human Controls and Risk", "Where automation must stop, ask, or escalate")
    guidance(doc, "Do not imply this section is a legal or cybersecurity audit. Record the workflow controls and flag specialist review when required.")
    table(doc, ["Risk or exception", "Preventive control", "Detection / evidence", "Owner and response"], [
        ("Incomplete or invalid input", "Required fields and validation rules.", "Rejected-item log and exception queue.", "Workflow owner reviews and corrects."),
        ("Sensitive data exposure", "Access limits and approved data boundaries.", "Access log and periodic review.", "System owner removes access and escalates."),
        ("Incorrect AI output", "Approved sources, prompt constraints, and human review.", "Sample audit and correction record.", "Approver corrects and reports pattern."),
        ("Failed integration", "Retry rules, idempotency, and safe defaults.", "Failure alert and transaction log.", "Automation owner resolves or reprocesses."),
        ("Material decision", "No autonomous approval or irreversible action.", "Approval record and named decision maker.", "Authorized approver accepts or rejects."),
    ], [2100, 2600, 2350, 2310])
    heading(doc, "Exception path")
    for item in (
        "Pause the automated path when required data, confidence, or approval is missing.",
        "Create a visible exception item with owner, reason, and due date.",
        "Preserve the input, output, decision, and correction needed for review.",
        "Use recurring exceptions to improve the process, rules, or training.",
    ):
        list_item(doc, item, numbered=True)

    page_break(doc)

    # 08 Roadmap
    section_title(doc, "08", "30-Day Roadmap", "A controlled path from decision to pilot")
    guidance(doc, "Assign a real owner and acceptance check to every action. Adjust timing for procurement, security review, or data preparation.")
    table(doc, ["Week", "Objective", "Actions", "Owner", "Acceptance check"], [
        ("1", "Define", "Confirm scope, owner, baseline, rules, exceptions, and approval boundary.", "Name owner", "Current state and success measure approved."),
        ("2", "Prepare", "Clean required data, document rules, confirm access, and configure a safe test environment.", "Name owner", "Inputs, permissions, and test cases ready."),
        ("3", "Pilot", "Build or configure the minimum useful workflow and test representative cases.", "Name owner", "Pilot completes cases without uncontrolled action."),
        ("4", "Measure", "Review results, errors, adoption, and controls; decide whether to expand.", "Name owner", "Decision made using recorded evidence."),
    ], [900, 1400, 3400, 1300, 2360])
    heading(doc, "Dependencies")
    for label, text in (
        ("People: ", "Owner, approvers, administrator, and pilot users."),
        ("Data: ", "Required records, clean-up, fields, retention, and access."),
        ("Technology: ", "Licenses, environments, integrations, and monitoring."),
        ("Governance: ", "Approval, privacy, security, vendor, and audit requirements."),
    ):
        list_item(doc, label + text, bold_start=label)

    page_break(doc)

    # 09 Metrics
    section_title(doc, "09", "Success Measures and Ownership", "How the future workflow will be managed")
    guidance(doc, "Use a small number of metrics the owner can actually review. Avoid invented targets when no baseline exists.")
    table(doc, ["Metric", "Baseline", "Pilot target / review rule", "Source", "Owner"], [
        ("Cycle time", "Current baseline", "Proposed target or compare after pilot", "System / sample", "Role"),
        ("Staff effort", "Current baseline", "Proposed target or compare after pilot", "System / sample", "Role"),
        ("Error / rework", "Current baseline", "Proposed target or compare after pilot", "Issue log", "Role"),
        ("Completion / response", "Current baseline", "Proposed target or compare after pilot", "Workflow system", "Role"),
        ("Exceptions", "Current count or none", "Review volume, cause, and resolution time", "Exception log", "Role"),
        ("Adoption", "Not applicable before pilot", "Usage and completion by intended users", "System activity", "Role"),
    ], [1800, 1550, 2600, 1650, 1760])
    heading(doc, "Operating cadence")
    table(doc, ["Cadence", "Participants", "Review", "Decision"], [
        ("Weekly during pilot", "Owner, approver, builder", "Volume, failures, exceptions, user feedback", "Correct, continue, pause, or expand"),
        ("Monthly after launch", "Owner and business sponsor", "Outcomes, controls, costs, and improvement backlog", "Prioritize next improvement"),
    ], [1700, 2200, 3260, 2200])

    page_break(doc)

    # 10 Recommendation
    section_title(doc, "10", "Recommendation and Next Decision", "What CRS365 recommends the client do next")
    guidance(doc, "Make a direct recommendation: self-directed action, preparation, an implementation sprint, or no automation.")
    callout(doc, "Recommended decision", "Proceed / proceed with conditions / prepare first / do not automate. Replace with the supported conclusion.", "E8F2ED", GREEN)
    heading(doc, "Required decisions")
    for item in (
        "Approve the workflow scope and named owner.",
        "Approve the human-control and exception boundaries.",
        "Approve the tool and data approach.",
        "Approve the pilot success measures and 30-day roadmap.",
    ):
        list_item(doc, item, numbered=True)
    heading(doc, "Implementation paths")
    table(doc, ["Path", "Best when", "Client owns", "CRS365 role"], [
        ("Self-directed", "The team has technical capacity and clear ownership.", "Build, test, train, monitor, and improve.", "Blueprint clarification only."),
        ("Implementation Sprint", "The team wants CRS365 to configure and launch the workflow.", "Access, decisions, testing, and adoption.", "Separate scope and fixed project quote."),
        ("Preparation phase", "Data, process, ownership, or governance is not ready.", "Documentation, cleanup, approvals, and readiness actions.", "Optional focused advisory support."),
    ], [1700, 2600, 2600, 2460])
    heading(doc, "Client acknowledgment")
    body(doc, "This Blueprint is an advisory deliverable based on information available during the engagement. It is not legal, cybersecurity, accounting, employment, or vendor warranty advice. Software capabilities and pricing can change. The client remains responsible for implementation decisions, access controls, testing, approvals, vendor terms, and operational outcomes.", size=9.5, shade_value=MUTED)
    body(doc, "Next review date: ____________________    Workflow owner: ____________________    Decision: ____________________", size=10.5, after=14)
    callout(doc, "CRS365", "Optimize. Automate. Operate.", SOFT, TEAL)

    doc.core_properties.title = "CRS365 AI Automation Blueprint Template"
    doc.core_properties.subject = "Fixed-scope AI automation advisory deliverable"
    doc.core_properties.author = "CRS365 Team"
    doc.core_properties.keywords = "AI automation, workflow, operations, roadmap, CRS365"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()